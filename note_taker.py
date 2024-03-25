# AI Notetaker

import argparse
import subprocess
import os
import sys
from pydub import AudioSegment
AudioSegment.converter = '/usr/local/bin/ffmpeg'
import openai
# import time
import tiktoken 
from time import time,sleep
import re
import docx
import datetime
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
# import shutil
import json

length = 0

max_length = 10000      # Adjust this depending on the number of tokens you have available.  3000 is the default.
max_tokens = 16000      # Max tokens we can use

gpt3_model = "gpt-3.5-turbo-0613"
gpt3_turbomodel = "gpt-3.5-turbo-16k-0613"



'''def gen_date_time_string() -> str:
    """
    This function generates and returns a string representation of the current date and time in the format
    YYYY.MM.DD-HH.MM.SS. The format is based on the format of a date and time string returned by the
    "date" command in Linux. This function is used when generating a log file name.
    """
    # get current time
    now = datetime.datetime.now()
    # create a string from the current time
    try:
        datestr = now.strftime("%Y.%m.%d-%H.%M.%S")
    except ValueError as err:
        print("Error: ", err)
    else:
        return str(datestr)
'''
## Initialize the error log, open the file.
'''def initialize_error_log(error_log):
    global error_log_file
    error_log_file = open(error_log, 'a', encoding='utf-8-sig')'''

def load_variables_from_file():
    '''
    Configuration variables are all set in the "meeting_config.key" file.  
    Loads up:
        - notes_folder_path, 
        - gpt_log_dir, and 
        - the api_key
        - errors_log path

    Returns a tuple of strings.
    '''
    # Variable python_file_dir is the directory the python file is in
    python_file_dir = os.path.dirname(os.path.abspath(__file__))

    with open(python_file_dir+'/meeting_notes_config.key', 'r') as file:

        def strip_quotes(string_in):
            # Strip quotes and new lines out of the string.
            string_out = string_in.replace('"', '')
            string_out = string_out.replace("'", '')
            string_out = string_out.replace("\n", '')
            return string_out

        config_file = file.read().splitlines()
        # Strip notes out of the config file, everything AFTER the # sign.
        config_file = [line.split('#', 1)[0] for line in config_file]
        # Strip out any blank lines.
        config_file = [line for line in config_file if line.strip()]
        # Strip out any lines that don't have an = sign.
        config_file = [line for line in config_file if "=" in line]

        # Now we have a list of lines that have a = sign in them.  Split them up into a dictionary.
        # Go through each line in the config file and split it into a dictionary.
        config_dict = {}
        for line in config_file:
            key, value = line.split("=", 1)
            config_dict[key.strip()] = value.strip()
        
        # Now we have a dictionary of key value pairs.  Load them into variables.
        notes_folder_path = config_dict["notes_folder_path"]
        api_key = config_dict["api_key"]
        gpt_log_dir = config_dict["gpt_log_dir"]
        notes_errors_log = config_dict["error_log"]

        api_key = strip_quotes(api_key)
        notes_folder_path = strip_quotes(notes_folder_path)
        gpt_log_dir = strip_quotes(gpt_log_dir)
        errors_log = strip_quotes(notes_errors_log)

        return notes_folder_path, api_key, gpt_log_dir, errors_log

### Shared Functions
def check_if_file_exists(file_path: str) -> bool:
    """
    Checks if a file exists.

    Args:
      file_path: The path to the file.

    Returns:
      True if the file exists, False otherwise.
    """
    return os.path.exists(file_path)

def create_directory(notes_path, file_name):
    """
    Create a directory if it does not exist.
    Parameters
        notes_path: The path to the notes directory
        file_name: The name of the directory to create
    Returns 
        directory_path: The path to the directory
    :example: create_directory("/home/notes", "notes_1") -> "/home/notes/notes_1"
    """
    directory_path = notes_path + "/" + file_name
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)
    return directory_path

def pydub_length(path):
    """
    Calculate the length of an audio file in milliseconds using the PyDub library.

    Args:
        path (str): The path to the audio file.

    Returns:
        int: The length of the audio file in milliseconds.

    Raises:
        Exception: If there is an error while processing the audio file.
    """
    try:
        audio = AudioSegment.from_mp3(path)
        length = len(audio)
        return length
    except Exception as e:
        print(f"Exception: {e}")
        return None

def convert_to_mp3(audio_file_path: str, bitrate: str = '128k') -> str:
    """
    Converts the audio file to a mono mp3 file with a specified bitrate.
    The bitrates for FFmpeg MP3 are 8, 16, 24, 32, 40, 48, 64, 80, 96, 112, 128, 160, 192, 224, 256, or 320.
    Args:
        audio_file_path: The path to the audio file.
        bitrate: The bitrate of the output mp3 file. Default is '128k'.

    Returns:
        A path to the mp3 file.

    Raises:
        ValueError: If the conversion failed.
    """
    mp3_file = 'output.mp3'

    if os.path.exists(mp3_file):
        os.remove(mp3_file)

    command = ['/usr/local/bin/ffmpeg', '-i', audio_file_path, '-ac', '1', '-b:a', bitrate, mp3_file]
    # command = ['/usr/local/Cellar/ffmpeg/5.1.2_1', '-i', audio_file_path, '-ac', '1', '-b:a', bitrate, mp3_file]
    subprocess.run(command, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    if not os.path.exists(mp3_file):
        raise ValueError('Conversion to mp3 failed.  Does file really exist?')

    return mp3_file

def convert_and_split_to_mp3(audio_file_path: str, output_folder: str = "output"):
    '''
    Convert the audio file to MP3 format and split it into smaller segments if necessary.

    Parameters:
    -----------
    audio_file_path : str
        The path to the input audio file.

    output_folder : str, optional
        The folder where the output MP3 files will be saved. Default is "output".

    Returns:
    --------
    mp3_files : list of str
        A list of paths to the MP3 files that were created. If the input audio file is smaller than 25 MB,
        the list will contain only one element, which is the path to the converted MP3 file. If the input
        audio file is larger than 25 MB, the list will contain multiple paths to the segmented MP3 files.

    Notes:
    ------
    - The function first converts the input audio file to MP3 format using a specified bitrate.
    - If the size of the resulting MP3 file is larger than 25 MB, the function splits it into 30-minute segments.
    - The segmented MP3 files are saved in the specified output folder.
    - The function returns a list of paths to the MP3 files that were created.

    Example:
    --------
    audio_file_path = "/path/to/audio.wav"
    output_folder = "output"
    mp3_files = convert_and_split_to_mp3(audio_file_path, output_folder)
    print(mp3_files)
    # Output: ['/path/to/output/part_1.mp3', '/path/to/output/part_2.mp3', ...]
    '''
    mp3_file = convert_to_mp3(audio_file_path, bitrate='48k')

    global length 
    length = pydub_length(mp3_file)
    # Load the MP3 into PyDub
    song = AudioSegment.from_mp3(mp3_file)

    # Check if the MP3 is over 25 MB
    if os.path.getsize(mp3_file) > 24 * 1024 * 1024:
        # Split the song into 30-min segments
        thirty_minutes = 20 * 60 * 1000
        segments = [song[i:i+thirty_minutes] for i in range(0, len(song), thirty_minutes)]

        # Create the output folder if it doesn't exist
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # Export the segments to the output folder
        mp3_files = []
        for idx, segment in enumerate(segments):
            segment_file = os.path.join(output_folder, f"part_{idx+1}.mp3")
            segment.export(segment_file, format="mp3")
            mp3_files.append(segment_file)

        return mp3_files
    else:
        return [mp3_file]

def count_tokens(post_text, token_model=gpt3_turbomodel):
    """
    Counts the number of tokens in the given post_text using the specified token_model.
    We use this to wrap a chat around a string, so we can run it through the model.  

    Args:
        post_text (str or list): The text or list of messages to count tokens from.
        token_model (str, optional): The token model to use for counting tokens. Defaults to gpt3_turbomodel.

    Returns:
        int: The number of tokens in the post_text.
    """
    if not isinstance(post_text, list):
        measure_message = [
            {"role": "system", "content": f"{post_text}"},
        ]
    else:
        measure_message = post_text

    number_of_words = num_tokens_from_messages(measure_message, model=token_model)
    return number_of_words

def num_tokens_from_messages(messages, model=gpt3_turbomodel):
    """
    Returns the number of tokens used by a list of messages.

    Args:
        messages (list): A list of messages.
        model (str, optional): The model to use for tokenization. Defaults to gpt3_turbomodel.

    Returns:
        int: The number of tokens used by the messages.
    """
    try:
        encoding = tiktoken.encoding_for_model(model)
    except KeyError:
        print("Warning: model not found. Using cl100k_base encoding.")
        encoding = tiktoken.get_encoding("cl100k_base")
    if model == "gpt-3.5-turbo":
        print("Warning: gpt-3.5-turbo may change over time. Returning num tokens assuming gpt-3.5-turbo-0301.")
        return num_tokens_from_messages(messages, model="gpt-3.5-turbo-0301")
    elif model == "gpt-4":
        print("Warning: gpt-4 may change over time. Returning num tokens assuming gpt-4-0314.")
        return num_tokens_from_messages(messages, model="gpt-4-0314")
    elif model == "gpt-3.5-turbo-0301":
        tokens_per_message = 4  # every message follows <|im_start|>{role/name}\n{content}<|end|>\n
        tokens_per_name = -1  # if there's a name, the role is omitted
    elif model == "gpt-4-0314":
        tokens_per_message = 3
        tokens_per_name = 1
    elif model == "gpt-3.5-turbo-16k-0613":
        tokens_per_message = 4
        tokens_per_name = -1
    else:
        raise NotImplementedError(f"""num_tokens_from_messages() is not implemented for model {model}. See https://github.com/openai/openai-python/blob/main/chatml.md for information on how messages are converted to tokens.""")
    num_tokens = 0
    for message in messages:
        num_tokens += tokens_per_message
        for key, value in message.items():
            num_tokens += len(encoding.encode(value))
            if key == "name":
                num_tokens += tokens_per_name
    num_tokens += 3  # every reply is primed with <|im_start|>assistant<|im_sep|>
    return num_tokens

def chat_with_gpt(messages, model="gpt-3.5-turbo", temperature=0.9, stop=[" Human:", " AI:"], presence_penalty=0.0, frequency_penalty=0.0, gpt_log_dir="gpt_logs/"):   
    """
    Chat with the GPT-3 model using OpenAI's Chat API.

    Args:
        messages (list): A list of message objects containing the conversation history.
        model (str, optional): The model to use for chat. Defaults to "gpt-3.5-turbo".
        temperature (float, optional): Controls the randomness of the output. Higher values make the output more random. Defaults to 0.9.
        stop (list, optional): A list of strings that, if encountered, will stop the response generation. Defaults to [" Human:", " AI:"].
        presence_penalty (float, optional): Controls the model's behavior to generate responses that are more or less verbose. Higher values make the output more verbose. Defaults to 0.0.
        frequency_penalty (float, optional): Controls the model's behavior to generate responses that are more or less repetitive. Higher values make the output less repetitive. Defaults to 0.0.
        gpt_log_dir (str, optional): The directory to store the GPT logs. Defaults to "gpt_logs/".

    Returns:
        str: The generated response from the GPT model.
    """
    max_retry = 5
    retry = 0
    while True:
        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=messages,
                temperature=temperature,
                stop=stop,
                presence_penalty=presence_penalty,
                frequency_penalty=frequency_penalty
            )

            filename = gpt_log_dir + '%s_gpt.log' % time()
            try:
                create_directory(gpt_log_dir, "")
            except:
                print("Error creating directory: " + gpt_log_dir)

            with open('%s' % filename, 'w') as outfile:
                try:
                    # assume we asked for a json output.  
                    outfile.write('PROMPT:\n\n' + str(json.dumps(messages, indent=1)) + '\n\n==========\n\nRESPONSE:\n\n' + str(response))
                except:
                    # assume we asked for a string output.
                    outfile.write('PROMPT:\n\n' + str(messages) + '\n\n==========\n\nRESPONSE:\n\n' + str(response))
                outfile.write(f'\n\nPARAMETERS:\n\n  model={model} \n\n temperature={str(temperature)} \n\n stop={str(stop)} \n\n presence_penalty={str(presence_penalty)} \n\n frequency_penalty={str(frequency_penalty)}')
            return response.choices[0].message.content
        except Exception as oops:
            retry += 1
            if retry >= max_retry:
                return "GPT3 error: %s" % oops
            print('Error communicating with OpenAI:' + str(oops))
            sleep(1)    #   3,000 Requests Per Minute is the limit for the API.  So, we'll wait 1 second between calls, just in case we ran into this?
            #               https://platform.openai.com/docs/guides/rate-limits/what-are-the-rate-limits-for-our-api

def consolidate_list_of_strings(list, max_length=2000):
    '''
    This function takes a list of strings and combines any two strings that are less than the max_length.
    Good practice to run this before running the GPT3 summarization function.  The larger text we feed GPT3,
    the better the results.

    Parameters
    ----------
    list : list
        A list of strings.
    max_length : int
        The maximum length of a string.  If two strings are less than this length, they will be combined.

    Returns
    -------
    list
        A list of strings.  Any two strings that were less than the max_length have been combined.

    '''
    paragraphs = list
    i = 0
    while i < len(paragraphs) - 1:
        # go through each paragraph. If the next paragraph will fit in the current paragraph, combine them.
        current_paragraph = paragraphs[i]
        next_paragraph = paragraphs[i + 1]
        if len(current_paragraph) + len(next_paragraph) + 1 < max_length:
            # if the next paragraph will fit in the current paragraph, combine them.
            paragraphs[i] = current_paragraph + " " + next_paragraph
            del paragraphs[i + 1]   # remove the next paragraph from the list.  We've combined it with the current paragraph.

        i += 1
    
    return paragraphs

def chunkify_text(text, max_length=2000, split_string="\n\n", debug_chunkify=False):
    ''' 
        This function takes a string text and a maximum length max_length (default 2000) 
        and returns a list of strings with max length max_length. It: 
        1. splits the text by newline characters into paragraphs, 
        2. consolidates paragraphs into the maximum length they can be
        3. then split each paragraph by sentences 
        and then check if adding a sentence to the current string will exceed the 
        max length or not. If it will, append the current string to the divided 
        text array and start a new string with the current sentence. If it won't, 
        add the sentence to the current string. Finally, it appends the last string 
        to the divided text array and returns the list.

        The goal is to get a body of text into the largest chunks possible, without
        breaking up sentences or paragraphs.  The larger the chunks, the better the 
        summary will be.
        Should see a gradual increase in the average string length of a paragraph,
        and very little change int he string lenght of the entire text:
        
        Total String Length 1: 70749
        Total String Length 2: 70529
        Avrg  String Length 2: 635
        Total String Length 3: 70580
        Avrg  String Length 3: 1176
        Total String Length 3: 70603
        Avrg  String Length 3: 1857

        Return a list of strings.
     '''

    # The goal is to get the largest chunks possible, without breaking up sentences or paragraphs.
    # So we're going to split the text by paragraphs.  
    # Then combine any paragraphs together that might fit under max_length.
    # Then split up any paragraphs that might be longer than max_length.
    
    if debug_chunkify: print("Total String Length 1: " + str(len(text)))

    # split the text by whatever we see as splitting up paragraphs. 
    # Default we're splitting them by "\n\n"
    paragraphs = text.split(split_string)
    
    if debug_chunkify:
        print("Total String Length 2: " + str(sum(len(string) for string in paragraphs)))  # Make sure we're not losing anything as we go through and chop it up.

    # Consolidate all the list of strings.  If there are any that are less than max_length, combine them.
    paragraphs = consolidate_list_of_strings(paragraphs, max_length=max_length)
    
    if debug_chunkify: 
        print("Total String Length 3: " + str(sum(len(string) for string in paragraphs)))  # Make sure we're not losing anything as we go through and chop it up.

    # Now go through paragraphs and if there are any that are longer than 2000, split them up.
    divided_text = []
    current_length = 0
    current_string = ""
    for paragraph in paragraphs:
        # split the paragraph by sentences
        sentences = re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', paragraph)
        for sentence in sentences:
            # check if adding the sentence will exceed the max length
            if current_length + len(sentence) > max_length:
                # if it will, add the current string to the divided text array
                divided_text.append(current_string)
                current_string = sentence
                current_length = len(sentence)
            else:
                # if it won't, add the sentence to the current string
                current_string += " " + sentence
                current_length += len(sentence)
    # append the last string to the divided text array
    divided_text.append(current_string)

    if debug_chunkify: 
        print("Total String Length 3: " + str(sum(len(string) for string in divided_text)))  # Make sure we're not losing anything as we go through and chop it up.

    return divided_text

def make_paragraphs(list_of_text_chunks):
    """
    Takes a list of text chunks as input and converts them into paragraphs.

    Parameters:
    - list_of_text_chunks (list): A list containing text chunks to be organized into paragraphs.
    
    Returns:
    - paragraphs (list): A list of paragraphs, where each paragraph is a string containing the organized and readable version of the original transcript.

    This function loops through each text chunk in the input list and generates a prompt message that explains the task to a language model AI. The AI model is called using the `chat_with_gpt` function, passing the prompt message and the text chunk as input. The AI generates an organized and readable version of the transcript and returns it as a string.

    The returned transcript is then split into a list of paragraphs using double newlines as separators, and each paragraph is appended to a list of paragraphs. 

    Finally, the function returns the list of paragraphs, where each paragraph is a string containing the organized and readable version of the original transcript.
    """
    
    count = 0

    paragraphs = [] # list of paragraphs.
    for text in list_of_text_chunks:
        # Print out our progress.
        count += 1
        print("Chunk " + str(count) + "  of " + str(len(list_of_text_chunks)))
        prompt = "As a helpful assistant, your task is to take a raw transcript of a meeting and improve its readability and organization by separating the text into logical paragraphs. Retain and Preserve all words and sentences in their original form when you write the new one. To separate each paragraph, use double newlines (a blank line between paragraphs).  Your goal is to provide an organized version of the original transcript that enables readers to easily read the original transcript from the meeting.  1. Remove all filler words and false starts. 2. Eliminate any sentences that are not significant or don't contribute to the overall meaning of the meeting. 3. Retain important details such as names, numbers, facts, and nouns. Please provide a revised version of the transcript that preserves essential information from the original.  Do not summarize or analyze the transcript."
        transcription_messages = [
            {"role": "system", "content": prompt},
            {"role": "assistant", "content": "Send me the transcript."},
            {"role": "user", "content": text},
            {"role": "assistant", "content": "What format would you like it returned in?"},
            {"role": "user", "content": "1. Remove all filler words and false starts. 2. Eliminate any sentences that are not significant or don't contribute to the overall meaning of the meeting. 3. Retain important details such as names, numbers, facts, and nouns. Please provide a revised version of the transcript that preserves essential information from the original.  Do not summarize or analyze the transcript."}
        ]
        answer = chat_with_gpt(transcription_messages, 
                               model=gpt3_turbomodel, 
                               frequency_penalty=0.125,
                               presence_penalty=0.125,
                               temperature=0.2)
        list_of_answers = answer.split("\n\n")
        paragraphs = paragraphs + list_of_answers
    return paragraphs

def transcribe(output_files, file_folder_path, logger, gpt_log_dir):
    """
    Transcribes the audio files in the given list and returns the full transcript.

    Output files is a list of files to be transcribed.  Returns a string of the discussiont.

    Go through each of the sound files in the list.  Open them, convert them to text, and save them to a file.  Returns the full transcript sewn together.
    Documentation [on transcription is here.](https://platform.openai.com/docs/api-reference/audio/create)

    Args:
        output_files (list): A list of files to be transcribed.
        file_folder_path (str): The path to the folder where the files are located.
        logger: The logger object used for logging.
        gpt_log_dir (str): The path to the directory where the raw transcription responses will be saved.

    Returns:
        str: The full transcript of the audio files.

    Raises:
        Exception: If an error occurs during transcription and the maximum number of retries is reached.
    """

    full_text_of_transcription = ""
    count_iter = 0 

    for file in output_files:
        # Open the mp3 audio file
        count_iter = count_iter + 1
        logger.debug("Transcribing: " + file)

        with open(file, "rb") as audio_file:
            # Transcribe the audio using the Whisper API
            
            max_retry = 5
            retry = 0
            while True:
                try:
                    transcription = openai.Audio.transcribe(file=audio_file,
                                                model="whisper-1", 
                                                response_format="json",
                                                temperature=0,
                                                language="en"
                                                )
            
                    break
                except Exception as oops:
                    retry += 1
                    if retry >= max_retry:
                        logger.debug("Transcribe error: %s" % oops)
                        quit()
                    logger.debug('Error transcribing:' + str(oops))
                    logger.debug("File: " + file)
                    logger.debug('Retrying...')
                    sleep(5)    #   3,000 Requests Per Minute is the limit for the API.  So, we'll wait 1 second between calls, just in case we ran into this?

        # save the raw response to file in the "gpt_logs" subfolder
        with open(f"{gpt_log_dir}/{file.split('/')[-1]}.json", "w") as file:
            file.write(json.dumps(transcription))
        # Print the transcription
        # logger.debug(transcription["text"])
        full_text_of_transcription += transcription["text"]

    logger.debug("Finished Transcribing.")
    logger.debug("Full Text of Transcription:" + full_text_of_transcription)

    # Write Transcription to File
    full_text_transcription_path = file_folder_path + "/full_text_transcription.txt"
    with open(full_text_transcription_path, "w") as file:
        file.write(full_text_of_transcription)
    
    return full_text_of_transcription

#######################
### Main Function

def main():
    
    ### Load up user defined variables
    notes_folder_path, api_key, gpt_log_dir, error_log_file = load_variables_from_file()
    logging.debug("Notes Folder Path: " + notes_folder_path)
    logging.debug("API Key: " + api_key)
    logging.debug("GPT Log Directory: " + gpt_log_dir)
    logging.debug("Error Log File: " + error_log_file)

    # load the openai key into the openai api
    openai.api_key = api_key

    logger = logging.getLogger(__name__)
    logging.basicConfig(filename=error_log_file, level=logging.DEBUG)
    stdout_handler = logging.StreamHandler(sys.stdout)
    stdout_handler.setLevel(logging.DEBUG)
    logger.addHandler(stdout_handler)

    logger.debug("Starting!")

    # This program takes in arguments from the command line.  If none are provided, it will ask for them.
    # Get the filename we're processing.  It can be taken in via command line, or via input after starting.  
    parser = argparse.ArgumentParser(description='AI Notetaker')
    parser.add_argument('--file', '-f', help='File to process', default=None)
    args = parser.parse_args()
    if args.file:
        if os.path.exists(args.file):
            logger.debug("File exists: " + args.file)
            original_file_path = args.file
        else:
            logger.debug("File does not exist: " + args.file)
            original_file_path = input("Paste your file path:  ")
    else:
        original_file_path = input("Paste your file path:  ")

    # Take the filepath and make it useable.  
    original_file_name = os.path.splitext(os.path.basename(original_file_path))[0]
    file_folder_path = create_directory(notes_folder_path, original_file_name)

    # Copy the file to the new directory, where we're working.  Check if it copied correctly.
    subprocess.call(["cp", original_file_path, file_folder_path])
    if check_if_file_exists(file_folder_path + "/" + original_file_name + ".m4a"):
        logger.debug("File copied successfully.")
    else:
        logger.debug("File not copied successfully.")
        logger.debug("Could not find file: " + file_folder_path + "/" + original_file_name + ".m4a")
        quit()  # don't proceed if we can't find the file.

    #!#!#!#!#!#!
    # Divide up the Audio.  Max audio size is 25 mb.  Stick it into the working folder.  Quit if error.
    file_path = file_folder_path + "/" + original_file_name + ".m4a"

    if not check_if_file_exists(file_path):
        raise ValueError(f"File {file_path} does not exist.")
        quit()
    output_files = convert_and_split_to_mp3(file_path)
    logger.debug(f"List of output files: {output_files}")

    audio_length = str(int(length/(60*1000))) + ':' + str(int((length/1000)%60))
    print(f"Meeting length: {audio_length} minutes.")

    full_text_of_transcription = transcribe(output_files, file_folder_path, logger, gpt_log_dir)
    paragraphs_in = chunkify_text(full_text_of_transcription, max_length=max_length, debug_chunkify=False)
    paragraphs_out = make_paragraphs(paragraphs_in)

    # Now we have a list of paragraphs.  Save it to text file.
    transcript_file_path = file_folder_path + "/transcript_" + original_file_name + ".txt"
    logger.debug("Saving transcript file: " + transcript_file_path)
    with open(transcript_file_path, "w") as f:
        for paragraph in paragraphs_out:
            f.write(paragraph + "\n\n")
    
    # Make a bullet point for each paragraph.
    bullet_points = [] # List of the bullet points
    iter_num = 0
    
    number_of_paras = len(paragraphs_out)
    for paragraph in paragraphs_out:
        iter_num = iter_num + 1
        logger.debug(f'Bulletizing {iter_num} of {number_of_paras} paragraphs.  {paragraph}')

        # Strip whitespace from the paragraph.  And if it's empty, skip it.
        paragraph = paragraph.lstrip().rstrip()
        if paragraph == "":
            logger.debug("Skipping empty paragraph.")
            continue

        # Check that the paragraph is at least 10 words long.  If not, skip it.
        if len(paragraph.split()) < 10:
            logger.debug("Skipping paragraph with less than 10 words.")
            continue

        transcription_messages = [
            {"role": "system", "content": "As an expert assistant in analyzing business conversations, your task is to provide a single bullet point summary of a paragraph from a meeting transcript. "},
            {"role": "assistant", "content": "Once I have received the paragraph from the meeting transcript, I will summarize the paragraph into a bullet point, ensuring it remains accurate and comprehensive in covering crucial elements from the conversation. "},
            {"role": "user", "content": "Please ensure your response focuses on accuracy and provides as many essential details as possible within these constraints while remaining thorough in its coverage of vital aspects from the conversation. "},
            {"role": "assistant", "content": "How should the output be formatted?"},
            {"role": "user", "content": "Return just the sentence of the bullet point summary.  Do not include any other information, label, number, or other characters other than the single sentence."},
            {"role": "assistant", "content": "Please provide the paragraph you'd like analyzed."},
            {"role": "user", "content": paragraph}
        ]

        bullet = chat_with_gpt(transcription_messages,
                            model=gpt3_turbomodel,
                            temperature=0.2,
                            frequency_penalty=0.25,
                            presence_penalty=0.25)

        bullet_points.append(bullet)
        logger.debug(f'Bulletized: {bullet}')

    # Save Bullet Points to Text File.      # Save the compressed transcript to a text file in the path compressed_transcript_file_path
    bullet_points_file_path = file_folder_path + "/bullets_" + original_file_name + ".txt"
    with open(bullet_points_file_path, "w") as f:
        for bullet_point in bullet_points:
            f.write(bullet_point + "\n\n")

    # Analyze the Text
    answer_list = []
    iter_num = 0
    compressed_transcript = "\n\n".join(paragraphs_out)

    # Check the length and number of words of the full transcript.  If you can squeeze it into 1/2 of 16k, analyze it in one shot.  If not, chunk it up.
    full_text_transcription_tokens = count_tokens(compressed_transcript, token_model=gpt3_turbomodel)
    if full_text_transcription_tokens > (max_tokens*1/2):
        logging.debug(f"Full text needs to be chunkified!  Full Text is {full_text_transcription_tokens} out of {max_tokens*1/2}")
        paragraphs_in = chunkify_text(compressed_transcript, max_length=max_length, debug_chunkify=False)
    else:
        logging.debug(f"Full text does NOT need to be chunkified!  Full Text is {full_text_transcription_tokens} out of {max_tokens*1/2}")
        paragraphs_in = [compressed_transcript]

    number_of_chunks = len(paragraphs_in)

    # [ ] Here's where there's a problem.  If we have a long meeting, we'll get multiple summaries.  We need to get this down to one summary.
    for chunk in paragraphs_in:
        iter_num = iter_num + 1
        logger.debug(f'Summarizing {iter_num} of {number_of_chunks}')

        transcription_messages = [
            {"role": "system", "content": "As an expert assistant in analyzing business conversations, your task is to provide a comprehensive summary and analysis of a meeting transcript. "},
            {"role": "assistant", "content": "Once I have received the required information (the meeting transcript), I will promptly begin working on your analysis, ensuring it remains accurate and comprehensive in covering crucial elements from the conversation.  What analysis do you want from the transcript?"},
            {"role": "user", "content": "Please ensure your response focuses on accuracy and provides as many essential details as possible within these constraints while remaining thorough in its coverage of vital aspects from the conversation. This analysis should be put into seven sections:  Your analysis should include: 1. A general meeting summary that consists of multiple paragraphs and ranges between 500-1500 words. 2. A list of action items, including deadlines or responsible parties. 3. A compilation of significant topics discussed during the meeting. 4. A record of questions asked and their corresponding answers. 5. Documentation of any unresolved questions from the discussion. 6. An overview of key decisions made throughout the meeting. 7. Up to five top keywords related to this meeting's content for search purposes."},
            {"role": "assistant", "content": "How should the output be formatted?"},
            {"role": "user", "content": "Separate out each section with double new lines.  Each section should have a title that starts with three #'s.  For example, the first section should be '# Meeting Summary'  The second section should be '# Action Items'  Do this for all seven sections."},
            {"role": "assistant", "content": "How should the analysis be structured?"},
            {"role": "user", "content": "Your analysis should include: 1. A general meeting summary that consists of multiple paragraphs and ranges between 500-1500 words. 2. A list of action items, including deadlines or responsible parties. 3. A compilation of significant topics discussed during the meeting. 4. A record of questions asked and their corresponding answers. 5. Documentation of any unresolved questions from the discussion. 6. An overview of key decisions made throughout the meeting. 7. Up to five top keywords related to this meeting's content for search purposes."},
            {"role": "assistant", "content": "Please provide the meeting transcript you'd like analyzed."},
            {"role": "user", "content": chunk}
        ]

        answer = chat_with_gpt(transcription_messages,
                            model=gpt3_turbomodel,
                            temperature=0.2,
                            frequency_penalty=0.125,
                            presence_penalty=0.125)

        answer_list.append(answer)

    # Save the Output to a text file.
    output_file_path = file_folder_path + "/analysis_" + original_file_name + ".txt"
    with open(output_file_path, "w") as f:
        # write each element of answer_list to file.
        for answer in answer_list:
            f.write(answer)
        
    word_doc_path = file_folder_path + "/Meeting-Notes-" + original_file_name + ".docx"

    # Open the text file up.  
    with open(output_file_path, "r") as file:
        lines = file.readlines()

    ###########################################################################
    # Step 2: Create a new Word document using python-docx
    doc = docx.Document()

    # Add Meta Data.  https://python-docx.readthedocs.io/en/latest/api/document.html#coreproperties-objects
    core_properties = doc.core_properties
    core_properties.author = 'John Cole'
    core_properties.title = f'Meeting Analysis and Notes: {original_file_name}'
    core_properties.subject = f'Notes'
    core_properties.category = f'Meeting Notes'

    # Set header information for all pages
    header = doc.sections[0].header
    header_text = f'Meeting Analysis and Notes: {original_file_name}    Written by John Cole.  Written On: {datetime.date.today()}'
    header.paragraphs[0].text = header_text

    # Add a title page.
    doc.add_heading(f'Meeting Title: {original_file_name}', level=0)
    doc.add_heading(f'Meeting Date: ', level=1)
    doc.add_heading('Written by:  John Cole', level=1)
    doc.add_heading(f'Attendees: ', level=1)
    doc.add_heading(f"Duration in minutes: {audio_length}")
    doc.add_page_break()

    # Add the long text transcrption to the end of the document.
    for line in lines:
        try:
            line = line.strip()
            # if line.startswith("###"):
            if line.startswith("###") or line.startswith("##") or line.startswith("#"):
                # Strip the leading ##'s off
                line = line.lstrip("#")
                doc.add_heading(line, level=2)

            elif line[0].isdigit() and line[1] == ".":
                doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(line)
        except Exception as e:
            logger.debug(f"Warning! Error in moving through analytical lines. {e}")

    ## Add in the bullet points of the meeting.
    doc.add_page_break()
    doc.add_heading("Detailed Bullet Points", level=1)
    
    logging.debug("Adding bullet points to the word doc.")
    for bullet in bullet_points:
        logging.debug("Adding Bullet: " + bullet)
        doc.add_paragraph(bullet, style="List Bullet")

    # Make a new page in the document.
    doc.add_page_break()
    doc.add_page_break()
    doc.add_heading("Cleaned Transcript", level=1)

    for paragraph in paragraphs_out:
        doc.add_paragraph(paragraph)

    # Make a new page in the document.
    doc.add_page_break()
    doc.add_page_break()
    doc.add_heading("Raw Transcript", level=1)

    doc.add_paragraph(full_text_of_transcription)

    # Step 4: Save the Word document
    doc.save(word_doc_path)

    # logger.debug("Finished writing to word doc.  Open here: " + file_name)
    logger.info(f"Finished writing to word doc.  Open here: \n \"{word_doc_path}\"")

    os.system('say "Finished processing file.  Please check!"')

if __name__ == "__main__":
    main()